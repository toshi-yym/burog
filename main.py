from fastapi import FastAPI, Request
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import requests
from bs4 import BeautifulSoup
from docx import Document
import tempfile
import os

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class UrlRequest(BaseModel):
    url: str

@app.post("/extract")
def extract_blog(req: UrlRequest):
    url = req.url
    try:
        res = requests.get(url, timeout=10)
        res.raise_for_status()
        soup = BeautifulSoup(res.text, "html.parser")
        # タイトル抽出
        title = soup.title.string if soup.title else "タイトルなし"
        # 本文抽出（シンプルな例: <article>タグ or <main>タグ or <body>）
        article = soup.find('article')
        if not article:
            article = soup.find('main')
        if not article:
            article = soup.body
        # 不要な要素の削除
        for selector in ['.profile', '.author', '.company', '.shop', '.service', '.footer', '.sidebar', '.related', '.recommend', '.nav', 'header', 'footer', 'aside']:
            for tag in article.select(selector):
                tag.decompose()
        # テキスト整形
        paragraphs = [p.get_text(strip=True) for p in article.find_all(['p', 'li', 'h2', 'h3', 'h4'])]
        body = "\n".join([p for p in paragraphs if p])
        # Wordファイル作成
        doc = Document()
        doc.add_heading(f'【{title}】', 0)
        for para in body.split('\n'):
            doc.add_paragraph(para)
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
        filename = "output.docx"
        return FileResponse(tmp_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=filename)
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=400)
